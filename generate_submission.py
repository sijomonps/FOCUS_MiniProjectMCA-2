
import sys
import subprocess
import os

# Try to import python-docx, install if missing
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    
    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title Page
    title = doc.add_heading('Mini Project Submission', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\nProject Title: FOCUS\n').bold = True
    p.add_run('A Personalized Learning Platform\n\n')
    p.add_run('Submitted by: [Your Name/ID]\n')
    p.add_run('Department of Computer Applications\n')
    p.add_run('Marian College Kuttikkanam (Autonomous)\n')
    
    doc.add_page_break()

    # 1. Project Proposal
    doc.add_heading('1. Project Proposal', level=1)
    
    doc.add_heading('Project Title', level=2)
    doc.add_paragraph('FOCUS: A Personalized Learning & Habit Building Platform')
    
    doc.add_heading('Problem Statement', level=2)
    doc.add_paragraph(
        "Students often struggle with procrastination, disorganized study schedules, and a lack of motivation. "
        "Existing tools are either too complex (project management software) or too simple (basic to-do lists), "
        "often failing to provide the visual feedback and 'habit-building' mechanics necessary for consistent academic progress. "
        "Furthermore, managing assignment deadlines alongside daily study sessions can be overwhelming without a unified interface."
    )
    
    doc.add_heading('Objectives', level=2)
    obj_p = doc.add_paragraph()
    obj_p.style = 'List Bullet'
    obj_p.add_run("To develop a distraction-free study timer that tracks focus sessions by subject.")
    doc.add_paragraph("To implement a visual 'Treemap' for assignments that highlights urgency intuitively.", style='List Bullet')
    doc.add_paragraph("To track and display daily study streaks to encourage consistent learning habits.", style='List Bullet')
    doc.add_paragraph("To provide a centralized dashboard for notes, tasks, and productivity analytics.", style='List Bullet')
    
    doc.add_heading('Proposed Solution', level=2)
    doc.add_paragraph(
        "FOCUS is a web-based application designed to bridge the gap between task management and habit tracking. "
        "It features a 'Focus Mode' timer, an intelligent assignment tracker that visually scales tasks based on deadlines, "
        "and a personal dashboard that greets the user with their daily progress. The application utilizes a dark, "
        "Notion-inspired aesthetic to minimize eye strain and maintain a calm study environment."
    )
    
    doc.add_heading('Technology Stack', level=2)
    stack_p = doc.add_paragraph()
    stack_p.add_run("Framework: ").bold = True
    stack_p.add_run("Django 6.0 (Python)\n")
    stack_p.add_run("Frontend: ").bold = True
    stack_p.add_run("HTML5, Vanilla CSS3, Vanilla JavaScript\n")
    stack_p.add_run("Database: ").bold = True
    stack_p.add_run("SQLite (Development)\n")
    stack_p.add_run("Visualization: ").bold = True
    stack_p.add_run("Chart.js")
    
    doc.add_heading('Expected Outcomes', level=2)
    doc.add_paragraph("By the end of this project, we will have a fully functional web application that allows users to register, log in, track their study hours in real-time, visualize pending assignments by urgency, and review their historical study performance.")
    
    doc.add_page_break()

    # 2. ER Diagram
    doc.add_heading('2. Entity-Relationship (ER) Diagram', level=1)
    doc.add_paragraph("The following section describes the entities, attributes, and relationships modeled in the system.")
    
    # User
    doc.add_heading('Entity: User (Django Auth)', level=3)
    doc.add_paragraph("Attributes: ID (PK), Username, Password, Email, Date_Joined.")
    
    # StudySession
    doc.add_heading('Entity: StudySession', level=3)
    doc.add_paragraph("Represents a completed block of study time.")
    p = doc.add_paragraph()
    p.add_run("Attributes: ").bold = True
    p.add_run("ID (PK), Subject, Duration (int), Date, Created_At.")
    p.add_run("\nRelationships: ").bold = True
    p.add_run("Many-to-One with User.")

    # Assignment
    doc.add_heading('Entity: Assignment', level=3)
    doc.add_paragraph("Represents a task with a deadline.")
    p = doc.add_paragraph()
    p.add_run("Attributes: ").bold = True
    p.add_run("ID (PK), Title, Subject, Deadline, Estimated_Hours, Status (Pending/Completed), Urgency (Low/Med/High).")
    p.add_run("\nRelationships: ").bold = True
    p.add_run("Many-to-One with User.")

    # QuickNote
    doc.add_heading('Entity: QuickNote', level=3)
    doc.add_paragraph("Represents short reflections added after study sessions.")
    p = doc.add_paragraph()
    p.add_run("Attributes: ").bold = True
    p.add_run("ID (PK), Title, Content, Study_Duration.")
    p.add_run("\nRelationships: ").bold = True
    p.add_run("Many-to-One with User, Many-to-One with SubjectFolder.")
    
    # Diagram Placeholder
    doc.add_paragraph("\n[Check Diagram Below]\n").alignment = WD_ALIGN_PARAGRAPH.CENTER
    # We can't draw easily, but we can make a text box representation
    diagram_text = (
        "      +-------+                 +--------------+\n"
        "      | User  | 1 <-------- *   | StudySession |\n"
        "      +-------+                 +--------------+\n"
        "          ^ 1\n"
        "          |\n"
        "          v *\n"
        "   +------------+               +--------------+\n"
        "   | Assignment |               | QuickNote    |\n"
        "   +------------+               +--------------+\n"
        "                                       * ^\n"
        "                                         |\n"
        "                                       1 |\n"
        "                                +---------------+\n"
        "                                | SubjectFolder |\n"
        "                                +---------------+"
    )
    p = doc.add_paragraph(diagram_text)
    for run in p.runs:
        run.font.name = 'Courier New'

    doc.add_page_break()

    # 3. Project Structure
    doc.add_heading('3. Django Project Structure', level=1)
    doc.add_paragraph("The project was initialized using standard Django commands.")
    doc.add_paragraph("Command used: django-admin startproject studyflow .")
    doc.add_paragraph("App created: python manage.py startapp core")
    
    doc.add_heading('Directory Hierarchy', level=3)
    structure = (
        "studyflow/\n"
        "  ├── manage.py\n"
        "  ├── studyflow/ (Settings)\n"
        "  ├── core/ (Main App)\n"
        "  │   ├── models.py\n"
        "  │   ├── views.py\n"
        "  │   └── urls.py\n"
        "  ├── templates/\n"
        "  └── static/"
    )
    p = doc.add_paragraph(structure)
    for run in p.runs:
        run.font.name = 'Courier New'
    doc.add_paragraph("The server runs successfully on http://127.0.0.1:8000/.")

    doc.add_page_break()

    # 4. SRS
    doc.add_heading('4. Requirements Specification', level=1)
    
    doc.add_heading('Functional Requirements', level=2)
    doc.add_paragraph("1. Authentication: Users must be able to sign up and log in securely.")
    doc.add_paragraph("2. Timer: Users can start, pause, and stop a study timer, selecting a subject.")
    doc.add_paragraph("3. Assignments: Users can add assignments with deadlines and mark them as complete.")
    doc.add_paragraph("4. Analytics: The system must display a weekly study chart and calculate streaks.")
    
    doc.add_heading('Non-Functional Requirements', level=2)
    doc.add_paragraph("1. Usability: The interface must be clean, 'distraction-free', and intuitive.")
    doc.add_paragraph("2. Performance: Dashboard and charts should load within 2 seconds.")
    doc.add_paragraph("3. Responsiveness: The application must be usable on mobile devices.")
    doc.add_paragraph("4. Reliability: Study data must be strictly persisted in the SQLite database.")

    doc.add_heading('Assumptions & Constraints', level=2)
    doc.add_paragraph("Assumption: Users verify their own study time (honesty system).")
    doc.add_paragraph("Constraint: Initial version uses SQLite; migration to PostgreSQL recommended for production.")

    doc.add_page_break()
    
    # 5. Hello World
    doc.add_heading('5. Basic Django "Hello World"', level=1)
    doc.add_paragraph("Status: COMPLETED")
    doc.add_paragraph(
        "The development environment is successfully set up. "
        "The project runs using 'python manage.py runserver'. "
        "Instead of a basic 'Hello World', the root URL directs to the Login page, "
        "or the Dashboard if authenticated, confirming the installation works."
    )

    # Save
    filename = "FOCUS_Project_Submission.docx"
    doc.save(filename)
    print(f"Document saved as {filename}")

if __name__ == "__main__":
    create_document()
